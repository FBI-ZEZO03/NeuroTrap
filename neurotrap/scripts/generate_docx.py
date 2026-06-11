"""
Generate NeuroTrap CADN Graduation Project Documentation as a .docx file.
Follows Al-Zaytoonah University (ZUJ) GP formatting rules exactly.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = "/home/neurotrap/neurotrap/Documentation/NeuroTrap_CADN_GP_Documentation.docx"

# ─── helpers ──────────────────────────────────────────────────────────────────

def set_page_margins(doc):
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.right_margin  = Cm(3.0)
        section.left_margin   = Cm(2.0)

def set_normal_style(doc):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    # Set cs/hAnsi fonts via XML so non-Latin chars use Times New Roman too
    rpr = style.element.get_or_add_rPr()
    el = OxmlElement('w:rFonts')
    el.set(qn('w:ascii'), 'Times New Roman')
    el.set(qn('w:hAnsi'), 'Times New Roman')
    el.set(qn('w:cs'),    'Times New Roman')
    rpr.append(el)

def para(doc, text, bold=False, italic=False, size=12,
         align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6,
         line_spacing=WD_LINE_SPACING.SINGLE, keep_with_next=False,
         color=None):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.line_spacing_rule = line_spacing
    if keep_with_next:
        pf.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold   = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def chapter_heading(doc, text):
    """Phase title  e.g.  'PLANNING PHASE' """
    p = para(doc, text, bold=True, size=16,
             align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=12, space_after=12,
             line_spacing=WD_LINE_SPACING.ONE_POINT_FIVE)
    return p

def section_heading(doc, text):
    """Section heading   e.g.  '1. Introduction' """
    p = para(doc, text, bold=True, size=16,
             space_before=12, space_after=6,
             line_spacing=WD_LINE_SPACING.ONE_POINT_FIVE)
    return p

def subsection_heading(doc, text):
    """Sub-section   e.g.  '1-1 Background' """
    p = para(doc, text, bold=True, size=14,
             space_before=8, space_after=4,
             line_spacing=WD_LINE_SPACING.ONE_POINT_FIVE)
    return p

def subsubsection_heading(doc, text):
    """3rd level heading """
    p = para(doc, text, bold=True, size=12,
             space_before=6, space_after=3)
    return p

def body(doc, text, space_after=6):
    return para(doc, text, size=12, space_after=space_after)

def placeholder_fig(doc, fig_id, description, caption):
    """Figure placeholder box."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after  = Pt(4)
    run = p.add_run(
        f"[FIGURE {fig_id} PLACEHOLDER]\n{description}"
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = RGBColor(128, 0, 0)

    # Caption BELOW the figure (ZUJ rule)
    cap = para(doc, caption, italic=True, size=11,
               align=WD_ALIGN_PARAGRAPH.CENTER,
               space_before=2, space_after=12)
    return p

def table_caption(doc, caption):
    """Table title ABOVE the table (ZUJ rule)."""
    return para(doc, caption, bold=True, size=12,
                space_before=12, space_after=3)

def add_table(doc, headers, rows, style='Table Grid'):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = style
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
    # data rows
    for r_idx, row in enumerate(rows):
        tbl_row = tbl.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row):
            cell = tbl_row.cells[c_idx]
            cell.text = str(cell_text)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
    # space after table
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return tbl

def add_page_break(doc):
    doc.add_page_break()

def ruled_line(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),  'single')
    bottom.set(qn('w:sz'),   '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ─── document build ───────────────────────────────────────────────────────────

doc = Document()
set_page_margins(doc)
set_normal_style(doc)

# Remove default empty paragraph
for p in doc.paragraphs[:]:
    if not p.text:
        p._element.getparent().remove(p._element)
        break

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
para(doc, "AL-ZAYTOONAH UNIVERSITY OF JORDAN", bold=True, size=16,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=36, space_after=6)
para(doc, "Faculty of Science and Information Technology", bold=True, size=14,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=6)
para(doc, "Department of Cybersecurity", bold=True, size=14,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=30)

ruled_line(doc)

para(doc,
     "NeuroTrap CADN: A Cognitive Adaptive Deception Network\n"
     "for Intelligent Honeypot-Based Active Defense and Attacker Profiling",
     bold=True, size=16,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=24, space_after=24)

ruled_line(doc)

para(doc,
     "Graduation Project Submitted to the Department of Cybersecurity in Partial\n"
     "Fulfillment of the Requirements for the Bachelor's Degree in Cybersecurity",
     size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=18, space_after=24)

para(doc, "Prepared by:", bold=True, size=12,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6)
para(doc, "[PLACEHOLDER: Student 1 Full Name]  -  [ID Number]", size=12,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "[PLACEHOLDER: Student 2 Full Name]  -  [ID Number]", size=12,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "[PLACEHOLDER: Student 3 Full Name]  -  [ID Number]", size=12,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

para(doc, "Supervised by:", bold=True, size=12,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=6)
para(doc, "[PLACEHOLDER: Supervisor Full Name, Academic Title]", size=12,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

para(doc, "Academic Year: 2025/2026", bold=True, size=12,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=0)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# CERTIFICATION
# ══════════════════════════════════════════════════════════════════════════════
para(doc, "CERTIFICATION", bold=True, size=16,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=18)

body(doc,
     "We, the undersigned, certify that we have read this graduation project titled "
     "\"NeuroTrap CADN: A Cognitive Adaptive Deception Network for Intelligent "
     "Honeypot-Based Active Defense and Attacker Profiling\" and in our opinion it "
     "meets the required standard for submission in partial fulfillment of the "
     "requirements for the Bachelor's Degree in Cybersecurity at Al-Zaytoonah "
     "University of Jordan.", space_after=18)

for title, role in [
    ("[PLACEHOLDER: Supervisor Full Name, Academic Title]", "Supervisor"),
    ("[PLACEHOLDER: Department Head Full Name, Academic Title]", "Department Head"),
    ("[PLACEHOLDER: External Examiner Full Name, Academic Title]", "External Examiner"),
]:
    para(doc, role + ":", bold=True, size=12, space_before=12, space_after=3)
    body(doc, title, space_after=3)
    body(doc, "Signature: ________________________    Date: ________________",
         space_after=12)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# DEDICATION
# ══════════════════════════════════════════════════════════════════════════════
para(doc, "DEDICATION", bold=True, size=16,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=18)
para(doc,
     "[PLACEHOLDER: Insert dedication text here. Example: \"To our families, "
     "whose unwavering support made this work possible. To the cybersecurity "
     "community dedicated to building a safer digital world.\"]",
     italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# ACKNOWLEDGMENTS
# ══════════════════════════════════════════════════════════════════════════════
para(doc, "ACKNOWLEDGMENTS", bold=True, size=16,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=18)
body(doc,
     "[PLACEHOLDER: Insert personalized acknowledgment text here.]", space_after=6)
body(doc,
     "The project team would like to express sincere gratitude to "
     "[PLACEHOLDER: Supervisor Full Name] for invaluable guidance, continuous "
     "support, and expert advice throughout the development of this project. "
     "Appreciation is also extended to the Department of Cybersecurity at "
     "Al-Zaytoonah University of Jordan for providing the academic environment and "
     "resources that made this work possible. Thanks are also due to the "
     "open-source communities behind the Cowrie, OpenCanary, and Scapy projects, "
     "whose contributions to the field of network security formed the foundation "
     "of the honeypot layer of this system.", space_after=0)
add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# LIST OF ABBREVIATIONS
# ══════════════════════════════════════════════════════════════════════════════
para(doc, "LIST OF ABBREVIATIONS", bold=True, size=16,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=12)
table_caption(doc, "Table IV-1. List of Abbreviations")
abbr_rows = [
    ("ADT",         "Attacker Digital Twin"),
    ("API",         "Application Programming Interface"),
    ("CADN",        "Cognitive Adaptive Deception Network"),
    ("CBEE",        "Cognitive Bias Exploitation Engine"),
    ("C2",          "Command and Control"),
    ("CI/CD",       "Continuous Integration / Continuous Delivery"),
    ("DP",          "Differential Privacy"),
    ("ER",          "Entity Relationship"),
    ("ERD",         "Entity Relationship Diagram"),
    ("FedAvg",      "Federated Averaging"),
    ("FHIM",        "Federated Honeypot Intelligence Mesh"),
    ("GADCF",       "Generative Adaptive Deception Content Factory"),
    ("IOC",         "Indicator of Compromise"),
    ("IP",          "Internet Protocol"),
    ("IPS",         "Intrusion Prevention System"),
    ("JSON",        "JavaScript Object Notation"),
    ("JSONL",       "JavaScript Object Notation Lines"),
    ("JWT",         "JSON Web Token"),
    ("KPI",         "Key Performance Indicator"),
    ("LLM",         "Large Language Model"),
    ("MFA",         "Multi-Factor Authentication"),
    ("MITRE ATT&CK","MITRE Adversarial Tactics, Techniques, and Common Knowledge"),
    ("ML",          "Machine Learning"),
    ("NLP",         "Natural Language Processing"),
    ("RDP",         "Remote Desktop Protocol"),
    ("REST",        "Representational State Transfer"),
    ("SMB",         "Server Message Block"),
    ("SMTP",        "Simple Mail Transfer Protocol"),
    ("SNMP",        "Simple Network Management Protocol"),
    ("SOC",         "Security Operations Center"),
    ("SSH",         "Secure Shell"),
    ("SSL",         "Secure Sockets Layer"),
    ("SVM",         "Support Vector Machine"),
    ("TCP",         "Transmission Control Protocol"),
    ("TLS",         "Transport Layer Security"),
    ("TOTP",        "Time-Based One-Time Password"),
    ("TTP",         "Tactic, Technique, and Procedure"),
    ("UDP",         "User Datagram Protocol"),
    ("UFW",         "Uncomplicated Firewall"),
    ("UML",         "Unified Modeling Language"),
    ("VNC",         "Virtual Network Computing"),
    ("VPS",         "Virtual Private Server"),
    ("WebSocket",   "WebSocket Protocol"),
]
add_table(doc, ["Abbreviation", "Full Term"], abbr_rows)
add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# LIST OF FIGURES
# ══════════════════════════════════════════════════════════════════════════════
para(doc, "LIST OF FIGURES", bold=True, size=16,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=12)
figs = [
    ("Figure 10-1",  "Project Gantt Chart"),
    ("Figure 10-2",  "Project PERT Chart"),
    ("Figure A-1-3", "System Context Diagram (DFD Level 0)"),
    ("Figure A-1-4", "DFD Level 0 - Main System Processes"),
    ("Figure A-1-5", "DFD Level 1 - Behavioral Analysis Process Decomposition"),
    ("Figure A-4-1", "System Use Case Diagram"),
    ("Figure A-4-2", "Sequence Diagram - Attack Detection and Response Flow"),
    ("Figure A-4-3", "Activity Diagram - Threat Response Workflow"),
    ("Figure A-6-1", "Entity Relationship Diagram"),
    ("Figure A-6-4", "System Class Diagram"),
    ("Figure D-2-1", "Dashboard Main Navigation Menu"),
    ("Figure D-3-1", "Login Form Design"),
    ("Figure D-3-2", "Attacker Profile Detail Modal"),
    ("Figure D-3-3", "CBEE Cognitive Bias Scorer Form"),
    ("Figure D-4-1", "SOC Incident Report Template"),
    ("Figure I-2-1", "Dashboard Main View - KPI Overview"),
    ("Figure I-2-2", "Live Events Feed Panel"),
    ("Figure I-2-3", "Threat Actors Table"),
    ("Figure I-2-4", "Honeypots Status Panel"),
    ("Figure I-2-5", "CBEE Cognitive Bias Panel"),
    ("Figure I-2-6", "GADCF Asset Generation Panel"),
    ("Figure I-2-7", "Attacker Digital Twin Detail Panel"),
    ("Figure I-2-8", "FHIM Federated Learning Panel"),
    ("Figure I-2-9", "AI SOC Analyst Panel"),
    ("Figure I-2-10","MITRE ATT&CK Technique Frequency Heatmap"),
    ("Figure I-2-11","Attacker Geolocation Map"),
    ("Figure I-2-12","System Architecture Deployment Diagram"),
]
for fig_id, fig_title in figs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(14.5), 2)  # right-align for page#
    run1 = p.add_run(f"{fig_id}  {fig_title}")
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(12)
    run2 = p.add_run("\t[TBD]")
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# LIST OF TABLES
# ══════════════════════════════════════════════════════════════════════════════
para(doc, "LIST OF TABLES", bold=True, size=16,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=12)
tbls = [
    ("Table IV-1",  "List of Abbreviations"),
    ("Table 8-3",   "SWOT Analysis"),
    ("Table 9-4",   "Economic Feasibility - Cost Estimation"),
    ("Table A-2-1", "Functional Requirements"),
    ("Table A-2-2", "Non-Functional Requirements"),
    ("Table A-2-3", "Software Requirements"),
    ("Table A-2-4", "Hardware Requirements"),
    ("Table A-5-1", "Decision Table - Threat Response Matrix"),
    ("Table A-6-5", "Business Rules"),
    ("Table D-1-1", "System Component Design Summary"),
    ("Table I-1-1", "Container Deployment Summary"),
    ("Table I-4-1", "Non-Functional Requirements Verification"),
]
for tbl_id, tbl_title in tbls:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run1 = p.add_run(f"{tbl_id}  {tbl_title}")
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(12)
    run2 = p.add_run("\t[TBD]")
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════
para(doc, "TABLE OF CONTENTS", bold=True, size=16,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=12)
toc_items = [
    ("Certification",                          "II"),
    ("Dedication",                             "III"),
    ("Acknowledgments",                        "IV"),
    ("List of Abbreviations",                  "V"),
    ("List of Figures",                        "VI"),
    ("List of Tables",                         "VII"),
    ("Abstract (English)",                     "VIII"),
    ("Abstract (Arabic)",                      "IX"),
    ("PLANNING PHASE",                         ""),
    ("  1.  Introduction",                     "1"),
    ("  2.  Overview of the Project",          "3"),
    ("  3.  Importance of the Project",        "5"),
    ("  4.  Current Business Description",     "7"),
    ("  5.  Current System Difficulties",      "10"),
    ("  6.  Problem Definition",               "12"),
    ("  7.  Scope of the System",              "14"),
    ("  8.  Strategic Planning",               "16"),
    ("  9.  Feasibility Analysis",             "20"),
    ("  10. Implementation Plan",              "25"),
    ("ANALYSIS PHASE",                         ""),
    ("  1.  Introduction to Analysis",         "27"),
    ("  2.  System Requirements",              "29"),
    ("  3.  Modeling Tools",                   "38"),
    ("  4.  UML Diagrams",                     "44"),
    ("  5.  Requirements Structuring",         "50"),
    ("  6.  Conceptual Data Modeling",         "56"),
    ("DESIGN PHASE",                           ""),
    ("  1.  System Design",                    "65"),
    ("  2.  Menu Design",                      "72"),
    ("  3.  Forms Design",                     "74"),
    ("  4.  Reports Design",                   "77"),
    ("IMPLEMENTATION PHASE",                   ""),
    ("  1.  System Deployment",                "79"),
    ("  2.  System Snapshots",                 "86"),
    ("  3.  System Maintenance and Support",   "93"),
    ("  4.  Non-Functional Requirements",      "96"),
    ("Conclusion and Future Work",             "99"),
    ("References",                             "103"),
    ("Appendix A - Source Code Excerpts",      "A1"),
    ("Appendix B - API Reference",             "B1"),
    ("Appendix C - Installation Guide",        "C1"),
]
for label, pg in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    bold_it = bool(label.strip()) and label.strip().isupper()
    run = p.add_run(label)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold_it
    if pg:
        run2 = p.add_run(f"  {'.' * max(1, 60 - len(label))}  {pg}")
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)
add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# ENGLISH ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════
para(doc, "ABSTRACT", bold=True, size=16,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=18)
body(doc,
     "NeuroTrap CADN (Cognitive Adaptive Deception Network) is an active defense "
     "cybersecurity platform that combines high-interaction and low-interaction "
     "honeypots with machine learning-based behavioral analysis, cognitive "
     "psychology-driven deception, and federated threat intelligence. The system "
     "captures attacker activity across eleven network services including SSH, "
     "Telnet, FTP, HTTP, SMB, MySQL, MSSQL, RDP, VNC, and SNMP. A ten-layer "
     "processing pipeline normalizes events, classifies attacker intent into six "
     "behavioral categories, computes dynamic threat scores, and generates "
     "personalized deception environments tailored to each attacker profile. "
     "The Cognitive Bias Exploitation Engine (CBEE) models five psychological "
     "dimensions to inject contextually relevant deceptive content, while the "
     "Generative Adaptive Deception Content Factory (GADCF) produces realistic "
     "fake assets. Federated learning with differential privacy enables threat "
     "intelligence sharing across organizational boundaries without exposing raw "
     "data. An AI-powered SOC Analyst module automates triage and incident "
     "reporting. The system is deployed on a production VPS using Docker containers "
     "and provides a real-time web dashboard with WebSocket-driven live feeds, "
     "MITRE ATT&CK heatmaps, geolocation mapping, and an AI analyst chat interface.",
     space_after=0)
add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# ARABIC ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════
para(doc, "ABSTRACT (ARABIC)", bold=True, size=16,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=18)
para(doc,
     "[PLACEHOLDER: Insert Arabic abstract here. The abstract must summarize "
     "the project in no more than 15 typed lines, right-to-left text direction, "
     "using standard Arabic cybersecurity terminology.]",
     italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# ═══════════════════════  PLANNING PHASE  ════════════════════════════════════
# ══════════════════════════════════════════════════════════════════════════════
chapter_heading(doc, "PLANNING PHASE")

# ── Section 1 ─────────────────────────────────────────────────────────────────
section_heading(doc, "1.  Introduction")
subsection_heading(doc, "1-1  Background")
body(doc,
     "The rapid expansion of interconnected digital infrastructure has created a "
     "continuously growing attack surface for malicious actors. Organizations "
     "worldwide face persistent threats from automated scanning bots, credential "
     "harvesting attacks, malware deployment campaigns, and sophisticated advanced "
     "persistent threat (APT) groups. Traditional defensive cybersecurity measures, "
     "including firewalls, intrusion detection systems, and signature-based "
     "antivirus software, operate reactively and provide limited insight into "
     "attacker tactics, techniques, and procedures (TTPs).")
body(doc,
     "Honeypot technology offers a fundamentally different defensive posture by "
     "deploying deliberately exposed decoy systems that attract, capture, and "
     "analyze malicious activity without risk to production assets [1]. A honeypot "
     "is a security resource whose value lies entirely in being probed, attacked, "
     "or compromised. The intelligence gathered from honeypot interactions enables "
     "defenders to understand attacker motivations, toolsets, and behavioral "
     "patterns at a level of detail unavailable through passive perimeter monitoring "
     "alone.")
body(doc,
     "NeuroTrap CADN addresses the limitations of existing static and isolated "
     "honeypot deployments by creating a cognitive, adaptive deception network "
     "that evolves in response to observed attacker behavior.")

subsection_heading(doc, "1-2  Motivation")
body(doc,
     "The motivation for this project stems from three converging observations. "
     "First, the volume and sophistication of cyberattacks continues to increase, "
     "with automated tools enabling low-skill actors to conduct large-scale "
     "intrusion campaigns. Second, existing open-source honeypot tools, while "
     "effective individually, lack integration with behavioral analytics and "
     "adaptive response capabilities. Third, cognitive psychology offers "
     "underutilized insights into attacker decision-making that can be exploited "
     "to extend attacker engagement and gather richer intelligence.")

subsection_heading(doc, "1-3  Project Objectives")
body(doc, "The primary objectives of the NeuroTrap CADN project are:")
objectives = [
    "To deploy and integrate multiple honeypot services covering eleven network protocols.",
    "To build an ML pipeline classifying attacker intent and assigning dynamic threat scores.",
    "To implement a Cognitive Bias Exploitation Engine (CBEE) exploiting attacker psychology.",
    "To develop a Generative Adaptive Deception Content Factory (GADCF) for fake asset production.",
    "To implement federated learning with differential privacy for cross-organizational intelligence sharing.",
    "To build an AI-powered SOC Analyst that automates triage and incident reporting.",
    "To provide a real-time web dashboard with comprehensive visualization of all subsystems.",
]
for i, obj in enumerate(objectives, 1):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(obj)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(3)

# ── Section 2 ─────────────────────────────────────────────────────────────────
section_heading(doc, "2.  Overview of the Project")
subsection_heading(doc, "2-1  System Description")
body(doc,
     "NeuroTrap CADN is a multi-layer active defense platform implemented as a "
     "collection of microservices deployed using Docker container orchestration. "
     "The system exposes deliberate attack surfaces across eleven network "
     "protocols and processes all interactions through a ten-layer analytical "
     "pipeline progressing from raw packet capture to automated threat response.")
body(doc,
     "The system runs on a production VPS with Ubuntu 24.04, six virtual CPUs, "
     "eleven gigabytes of RAM, and 193 gigabytes of SSD storage. The management "
     "dashboard is accessible via HTTPS on port 443.")

subsection_heading(doc, "2-2  Ten-Layer Processing Pipeline")
table_caption(doc, "Table 2-1. Ten-Layer Processing Pipeline")
pipeline_rows = [
    ("1",  "Capture",               "Cowrie JSONL tail and Scapy packet capture"),
    ("2",  "Detection",             "AlertEvent normalization and severity assignment"),
    ("3",  "Behavior Analysis",     "13-feature extraction, ML classification, threat scoring"),
    ("4",  "Deception Engine",      "Environment template selection and spawning"),
    ("5",  "CBEE",                  "Cognitive bias scoring and bait injection"),
    ("6",  "GADCF",                 "Generative deception content production"),
    ("7",  "Attacker Digital Twin", "Behavioral synthesis and attack prediction"),
    ("8",  "FHIM",                  "Federated learning aggregation and intelligence sharing"),
    ("9",  "Response Engine",       "Automated threat response actions"),
    ("10", "API and Dashboard",     "REST API, WebSocket streaming, visualization"),
]
add_table(doc, ["Layer", "Name", "Responsibility"], pipeline_rows)

# ── Section 3 ─────────────────────────────────────────────────────────────────
section_heading(doc, "3.  Importance of the Project")
subsection_heading(doc, "3-1  Academic Importance")
body(doc,
     "This project contributes to the academic field of active cyber defense by "
     "demonstrating the practical integration of network security, machine learning, "
     "natural language processing, cognitive psychology, and distributed systems. "
     "The project implements federated learning [9] and differential privacy [10] "
     "in a security context and applies the MITRE ATT&CK framework [2] for "
     "behavioral classification at the operational level.")
subsection_heading(doc, "3-2  Practical Importance")
body(doc,
     "NeuroTrap CADN demonstrates that a comprehensive threat intelligence platform "
     "can be built from open-source components at zero software licensing cost and "
     "deployed on modest hardware. The automated SOC analyst functionality reduces "
     "the human expertise required to interpret threat data.")
subsection_heading(doc, "3-3  Contribution to the Cybersecurity Community")
body(doc,
     "The project addresses three specific gaps: (1) static environments that "
     "allow attacker fingerprinting; (2) isolated operation without cross-"
     "organizational intelligence sharing; and (3) passive observation without "
     "attempts to influence attacker behavior using cognitive psychology principles.")

# ── Section 4 ─────────────────────────────────────────────────────────────────
section_heading(doc, "4.  Current Business Description")
subsection_heading(doc, "4-1  Overview of Existing Honeypot Technology")
body(doc,
     "The current landscape of honeypot and deception technology spans a spectrum "
     "from low-interaction to high-interaction systems [1]. Low-interaction "
     "honeypots emulate only network interfaces and basic protocol behaviors. "
     "High-interaction honeypots expose real or fully emulated operating system "
     "environments capturing richer behavioral data.")
body(doc,
     "The most widely deployed open-source honeypots include Cowrie [3] "
     "(SSH/Telnet, medium-to-high interaction), OpenCanary [4] (multi-protocol "
     "low-interaction), and Galah [5] (LLM-powered web honeypot). Dionaea "
     "(multi-protocol malware capture) is excluded because it triggers a SIGTRAP "
     "crash on Linux kernel 6.8 due to an incompatibility in the libemu library.")
subsection_heading(doc, "4-2  Current Operational Practice")
body(doc,
     "In current practice, threat intelligence from honeypot deployments is "
     "collected manually, analyzed by security analysts, and acted upon through "
     "manual firewall rule updates and incident reports. This process is slow, "
     "requires significant analyst expertise, and does not scale to the volume of "
     "automated attack traffic generated by modern botnets.")
subsection_heading(doc, "4-3  Standards Context")
body(doc,
     "The MITRE ATT&CK framework [2] provides the standard taxonomy for adversary "
     "tactics and techniques used throughout this project. Differential privacy "
     "parameters follow the mathematical framework of Dwork et al. [10].")

# ── Section 5 ─────────────────────────────────────────────────────────────────
section_heading(doc, "5.  Current System Difficulties")
for heading, text in [
    ("5-1  Static Honeypot Environments",
     "Static honeypot environments present the same decoy interface to every "
     "attacker regardless of their sophistication. This uniformity allows "
     "experienced attackers to fingerprint and identify honeypots, leading to "
     "early disengagement."),
    ("5-2  Absence of Behavioral Analysis Integration",
     "Existing tools record interaction data but do not natively analyze attacker "
     "behavior to classify intent, assess threat level, or predict future actions. "
     "Security analysts must manually review session logs, which is time-consuming "
     "and impossible to scale."),
    ("5-3  Isolated Intelligence Silos",
     "Each honeypot deployment operates as an independent intelligence source. "
     "Threat indicators captured by one organization are not automatically shared "
     "with others, limiting the collective defensive value of distributed networks."),
    ("5-4  High Analyst Workload",
     "A single public-IP honeypot may receive hundreds of connection attempts per "
     "day. Existing tools provide insufficient automation for log review, profile "
     "correlation, incident report generation, and firewall management."),
    ("5-5  Absence of Psychological Engagement Mechanisms",
     "No existing open-source honeypot framework implements cognitive psychology "
     "principles to actively influence attacker behavior and extend engagement "
     "duration."),
]:
    subsection_heading(doc, heading)
    body(doc, text)

# ── Section 6 ─────────────────────────────────────────────────────────────────
section_heading(doc, "6.  Problem Definition")
subsection_heading(doc, "6-1  Problem Statement")
body(doc,
     "The core problem is the inadequacy of passive, static, and isolated honeypot "
     "deployments for generating actionable threat intelligence. There is no "
     "integrated, open-source platform combining multi-protocol honeypot capture "
     "with machine learning behavioral analysis, adaptive deception environment "
     "generation, cognitive psychology-based engagement, and privacy-preserving "
     "federated intelligence sharing under a real-time operational dashboard.")
subsection_heading(doc, "6-2  List of Problems")
problems = [
    ("Problem 1 - Static Deception Environments",
     "Identical environments for all attackers enabling fingerprinting."),
    ("Problem 2 - No Automated Behavioral Classification",
     "No automated distinction between low-threat scanners and advanced human operators."),
    ("Problem 3 - No Threat Scoring Mechanism",
     "No standard method integrating ML confidence, TTP coverage, persistence, "
     "and sophistication into a single actionable score."),
    ("Problem 4 - No Cognitive Engagement Layer",
     "No exploitation of attacker cognitive biases to extend engagement duration."),
    ("Problem 5 - Intelligence Isolation",
     "No privacy-preserving mechanism to aggregate intelligence across organizations."),
    ("Problem 6 - High Analyst Workload",
     "No automated triage, report generation, or threat queuing."),
    ("Problem 7 - Incomplete Protocol Coverage",
     "No single platform covers SSH, Telnet, FTP, HTTP, SMB, MySQL, MSSQL, "
     "RDP, VNC, SNMP, and LLM HTTP in a unified framework."),
]
for title, desc in problems:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run1 = p.add_run(title + ": ")
    run1.bold = True
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(12)
    run2 = p.add_run(desc)
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)

# ── Section 7 ─────────────────────────────────────────────────────────────────
section_heading(doc, "7.  Scope of the System")
subsection_heading(doc, "7-1  In-Scope Functionality")
scope_items = [
    "Deployment and management of Cowrie, OpenCanary, and Galah honeypot services.",
    "Event capture and normalization into a unified AlertEvent schema.",
    "Per-session Cowrie event aggregation into complete session documents.",
    "13-dimensional behavioral feature extraction per attacker IP address.",
    "ML-based attacker intent classification (6 classes) and tier classification (3 classes).",
    "Composite threat score computation (0-100) integrating ML confidence, TTPs, persistence, and volume.",
    "Adaptive deception environment spawning from three templates.",
    "Cognitive bias profiling (5 dimensions) and personalized bait injection.",
    "Generation of five types of fake deception assets.",
    "Attacker Digital Twin with tactic prediction and kill-chain mapping.",
    "Federated learning with differential privacy for intelligence sharing.",
    "Four-tier automated response action execution.",
    "AI SOC Analyst for triage, incident reports, and Q&A.",
    "Real-time dashboard with WebSocket streaming, MITRE heatmap, and geo map.",
]
for item in scope_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(3)

subsection_heading(doc, "7-2  Out-of-Scope Items")
out_items = [
    "ASHRTA module (planned but not implemented in this version).",
    "Production load testing under high-volume attack traffic.",
    "Windows-specific honeypots beyond SMB emulation.",
    "Active offensive countermeasures against attacker infrastructure.",
    "Native mobile application for the SOC analyst interface.",
]
for item in out_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(3)

# ── Section 8 ─────────────────────────────────────────────────────────────────
section_heading(doc, "8.  Strategic Planning")
subsection_heading(doc, "8-1  System Constraints")
constraints = [
    ("Hardware",
     "Minimum four-core CPU, eight GB RAM, fifty GB disk. Production: 6 vCPU, 11 GB RAM, 193 GB SSD."),
    ("Network",
     "Public IP required. Real SSH management on port 50402; port 22 is owned by Cowrie."),
    ("Kernel Compatibility",
     "Dionaea excluded due to SIGTRAP crash on Linux kernel 6.8 (libemu incompatibility)."),
    ("LLM API",
     "Galah and LLM SOC reports require an API key. Without a key, Galah is disabled and SOC Analyst runs in heuristic mode."),
    ("Container Privilege",
     "Only packet-monitor has NET_ADMIN/NET_RAW capabilities. No container runs in privileged mode."),
    ("Privacy",
     "FHIM applies Gaussian differential privacy (epsilon=1.0, delta=1e-5) to all model deltas."),
]
for title, desc in constraints:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run1 = p.add_run(title + " Constraint: ")
    run1.bold = True; run1.font.name = 'Times New Roman'; run1.font.size = Pt(12)
    run2 = p.add_run(desc)
    run2.font.name = 'Times New Roman'; run2.font.size = Pt(12)

subsection_heading(doc, "8-2  Project Scope Boundaries")
body(doc,
     "Input boundary: all network traffic on honeypot ports and Cowrie JSONL events. "
     "Processing boundary: all analysis runs within Docker containers on the deployment server. "
     "Output boundary: MongoDB records, iptables rules, traffic shaping commands, "
     "alert notifications, and dashboard visualizations. "
     "External integration boundary: LLM APIs (Anthropic, OpenAI, Groq) and alert delivery services.")

subsection_heading(doc, "8-3  SWOT Analysis")
table_caption(doc, "Table 8-3. SWOT Analysis")
swot_rows = [
    ("Strengths",
     "1. Comprehensive protocol coverage (11 services).\n"
     "2. Integration of ML, NLP, and cognitive psychology.\n"
     "3. Real-time WebSocket-driven dashboard.\n"
     "4. Privacy-preserving federated learning.\n"
     "5. Fully containerized, reproducible deployment.\n"
     "6. Zero software licensing cost."),
    ("Weaknesses",
     "1. LLM-dependent features require external API keys.\n"
     "2. ASHRTA module is planned but not implemented.\n"
     "3. No formal load testing completed.\n"
     "4. Federated learning demo uses synthetic deltas only.\n"
     "5. Self-signed SSL certificate."),
    ("Opportunities",
     "1. Growing demand for affordable threat intelligence platforms.\n"
     "2. Local LLMs (Ollama/Mistral) enable fully offline operation.\n"
     "3. MITRE ATT&CK provides extensible classification vocabulary.\n"
     "4. Potential for academic research publication."),
    ("Threats",
     "1. Sophisticated attackers may detect and exit honeypot environments.\n"
     "2. LLM API rate limits and costs may constrain operation.\n"
     "3. Open-source honeypot signatures are publicly documented.\n"
     "4. Kernel/library changes may break components (as with Dionaea)."),
]
add_table(doc, ["Category", "Items"], swot_rows)

# ── Section 9 ─────────────────────────────────────────────────────────────────
section_heading(doc, "9.  Feasibility Analysis")
subsection_heading(doc, "9-1  Technical Feasibility")
body(doc,
     "Technical feasibility is confirmed by the successful implementation and "
     "deployment of all core modules on the production VPS. All required "
     "dependencies are available as open-source packages: Scapy [6] for packet "
     "capture, scikit-learn [7] for ML classification, sentence-transformers [8] "
     "for NLP-based TTP matching, Flask [12] for the API, and PyMongo [15] for "
     "database access.")
subsection_heading(doc, "9-2  Schedule Feasibility")
body(doc,
     "[PLACEHOLDER: Insert actual project timeline dates.] "
     "The project was organized into seven sequential phases: Infrastructure, "
     "Detection Pipeline, Behavioral Analysis, Deception and Response, "
     "Intelligence Modules, Dashboard and API, Testing and Hardening. "
     "All phases were completed within the 2025/2026 academic year.")
subsection_heading(doc, "9-3  Operational Feasibility")
body(doc,
     "All containers are configured with restart policy 'unless-stopped' for "
     "continuous unattended operation. The AI SOC Analyst generates daily shift "
     "summaries and maintains a prioritized triage queue, significantly reducing "
     "analytical burden on human operators.")
subsection_heading(doc, "9-4  Economic Feasibility")
table_caption(doc, "Table 9-4. Economic Feasibility - Cost Estimation")
econ_rows = [
    ("VPS hosting (6 vCPU, 11 GB RAM)",  "Recurring", "[PLACEHOLDER: Monthly USD]"),
    ("Cowrie honeypot",                   "Open source", "$0"),
    ("OpenCanary honeypot",               "Open source", "$0"),
    ("Galah web honeypot",                "Open source", "$0"),
    ("Scapy / scikit-learn / Flask / MongoDB / Docker", "Open source", "$0"),
    ("LLM API (Groq free tier)",          "External service", "$0"),
    ("LLM API (Anthropic, optional)",     "External service", "Usage-based"),
    ("Total software cost",               "",             "$0"),
    ("Total infrastructure cost",         "",             "[PLACEHOLDER]"),
]
add_table(doc, ["Item", "Type", "Cost"], econ_rows)

# ── Section 10 ────────────────────────────────────────────────────────────────
section_heading(doc, "10.  Implementation Plan")
subsection_heading(doc, "10-1  Project Gantt Chart")
placeholder_fig(doc, "10-1",
    "Insert the project Gantt chart showing all development phases, "
    "durations, and dependencies. Rows: Requirements Analysis, Infrastructure Setup, "
    "Detection Pipeline, Behavioral Analysis, Deception Modules, Intelligence Modules, "
    "Dashboard Development, Testing and Hardening, Documentation.",
    "Figure 10-1. Project Development Gantt Chart")
body(doc,
     "As illustrated in Figure 10-1, the project was executed in seven sequential "
     "phases over the 2025/2026 academic year.")

subsection_heading(doc, "10-2  Project PERT Chart")
placeholder_fig(doc, "10-2",
    "Insert the PERT network diagram showing task dependencies, critical path, "
    "earliest start times, latest finish times, and slack for each development activity.",
    "Figure 10-2. Project PERT Network Diagram")
body(doc,
     "The critical path analysis shown in Figure 10-2 identifies the behavioral "
     "analysis module and dashboard integration as the activities with zero slack.")


# ══════════════════════════════════════════════════════════════════════════════
# ═══════════════════════  ANALYSIS PHASE  ════════════════════════════════════
# ══════════════════════════════════════════════════════════════════════════════
add_page_break(doc)
chapter_heading(doc, "ANALYSIS PHASE")

# ── Section 1 ─────────────────────────────────────────────────────────────────
section_heading(doc, "1.  Introduction to Analysis")
subsection_heading(doc, "1-1  Introduction")
body(doc,
     "The Analysis Phase translates high-level project objectives into precise, "
     "technically grounded system requirements. This phase documents what the "
     "system must do (functional requirements), how well it must perform "
     "(non-functional requirements), what tools and hardware are needed, and how "
     "system processes relate to each other through modeling and UML diagrams.")
subsection_heading(doc, "1-2  System Methodology")
body(doc,
     "The system follows a microservices architecture pattern where each major "
     "functional component runs as an independent Docker container with well-defined "
     "interfaces through the shared MongoDB database and the Flask REST API. This "
     "methodology provides isolation, independent scaling, reproducibility, and "
     "maintainability.")

# ── Section 2 ─────────────────────────────────────────────────────────────────
section_heading(doc, "2.  System Requirements")
subsection_heading(doc, "2-1  Functional Requirements")
table_caption(doc, "Table A-2-1. Functional Requirements")
fr_rows = [
    ("FR-01","Capture all TCP/UDP connections on honeypot ports via Scapy.","High"),
    ("FR-02","Ingest Cowrie JSONL log events in real time.","High"),
    ("FR-03","Normalize all events into a unified AlertEvent schema.","High"),
    ("FR-04","Aggregate per-session Cowrie events into session documents.","High"),
    ("FR-05","Detect port scanning: more than 10 distinct ports in 5 seconds.","High"),
    ("FR-06","Detect brute force: more than 5 failed logins in 60 seconds.","High"),
    ("FR-07","Classify attacker intent into 6 categories.","High"),
    ("FR-08","Classify attacker tier into 3 categories.","High"),
    ("FR-09","Extract MITRE ATT&CK technique IDs from commands.","High"),
    ("FR-10","Compute composite threat score 0-100 per attacker.","High"),
    ("FR-11","Spawn deception environment when threat score >= 10.","High"),
    ("FR-12","Inject bait when bias overall >= 25 and injections < 3.","High"),
    ("FR-13","Generate 5 types of fake deception assets.","Medium"),
    ("FR-14","Build Attacker Digital Twin for each profiled attacker.","Medium"),
    ("FR-15","Predict next 3 MITRE tactics via Markov chain.","Medium"),
    ("FR-16","Support federated model sharing with differential privacy.","Medium"),
    ("FR-17","Execute block/isolate/slow/log response based on score.","High"),
    ("FR-18","Send alerts via SMTP, Slack, and Telegram.","Medium"),
    ("FR-19","Provide REST API with JWT authentication.","High"),
    ("FR-20","Support TOTP multi-factor authentication for admin.","Medium"),
    ("FR-21","Stream live events via WebSocket.","High"),
    ("FR-22","Provide AI SOC Analyst with triage and report generation.","Medium"),
    ("FR-23","Display MITRE ATT&CK technique frequency heatmap.","Low"),
    ("FR-24","Display attacker geolocation on interactive world map.","Low"),
    ("FR-25","Maintain persistent response_log for all response actions.","High"),
]
add_table(doc, ["ID", "Requirement", "Priority"], fr_rows)

subsection_heading(doc, "2-2  Non-Functional Requirements")
table_caption(doc, "Table A-2-2. Non-Functional Requirements")
nfr_rows = [
    ("NFR-01","Event-to-dashboard latency","Less than 5 seconds"),
    ("NFR-02","ML classifier macro F1 score","Greater than 0.85"),
    ("NFR-03","Deception environment spawn time","Less than 30 seconds"),
    ("NFR-04","Response action execution time","Less than 10 seconds"),
    ("NFR-05","Lynis host hardening score","Greater than 70"),
    ("NFR-06","Maximum simultaneous active environments","20 environments"),
    ("NFR-07","Environment maximum lifetime","60 minutes"),
    ("NFR-08","MongoDB availability","99%+ (restart: unless-stopped)"),
    ("NFR-09","API response cache duration","30 seconds"),
    ("NFR-10","JWT token expiry","1 hour"),
    ("NFR-11","FHIM differential privacy epsilon","1.0"),
    ("NFR-12","FHIM differential privacy delta","1e-5"),
]
add_table(doc, ["ID", "Requirement", "Target Value"], nfr_rows)

subsection_heading(doc, "2-3  Software Requirements")
table_caption(doc, "Table A-2-3. Software Requirements")
sw_rows = [
    ("Docker 24+",              "Container runtime and orchestration"),
    ("Cowrie (latest)",         "SSH and Telnet honeypot"),
    ("OpenCanary (latest)",     "Multi-protocol low-interaction honeypot"),
    ("Galah (latest)",          "LLM-powered web honeypot"),
    ("Scapy 2.5+ [6]",          "Raw packet capture and analysis"),
    ("scikit-learn 1.3+ [7]",   "VotingClassifier, RandomForest, SVM"),
    ("sentence-transformers [8]","TTP semantic embedding matching"),
    ("Flask 3.x [12]",          "REST API and web dashboard"),
    ("MongoDB 6.0 [13]",        "Document event and profile storage"),
    ("PyMongo 4.x [15]",        "Python MongoDB driver"),
    ("flask-jwt-extended [16]", "JWT token authentication"),
    ("pyotp [19]",              "TOTP multi-factor authentication"),
    ("Faker 20+ [18]",          "Fake credential and asset generation"),
    ("nginx Alpine",            "TLS termination and reverse proxy"),
    ("Anthropic API [17]",      "LLM SOC reports and Galah (optional)"),
    ("Ollama/Mistral [20]",     "Local LLM for GADCF (optional)"),
]
add_table(doc, ["Software", "Purpose"], sw_rows)

subsection_heading(doc, "2-4  Hardware Requirements")
table_caption(doc, "Table A-2-4. Hardware Requirements")
hw_rows = [
    ("CPU",      "4 cores",     "6+ cores",     "6 vCPU"),
    ("RAM",      "8 GB",        "12 GB",        "11 GB"),
    ("Disk",     "50 GB SSD",   "100 GB SSD",   "193 GB SSD"),
    ("Public IP","Required",    "Required",     "13.140.144.118"),
]
add_table(doc, ["Spec", "Minimum", "Recommended", "Production"], hw_rows)

subsection_heading(doc, "2-5  Development Environment")
body(doc,
     "The development environment uses the same Docker Compose configuration as "
     "production. A SQLite-backed FallbackDB class is available for local development "
     "without MongoDB, activated by setting NEUROTRAP_FORCE_FALLBACK=1. The CI/CD "
     "pipeline runs on GitHub Actions using Python 3.11 with pytest and ruff linting.")

# ── Section 3 ─────────────────────────────────────────────────────────────────
section_heading(doc, "3.  Modeling Tools")
subsection_heading(doc, "3-1  Process Modeling")
body(doc,
     "System processes are modeled using Data Flow Diagrams (DFD) at three levels "
     "of abstraction. External entities are: Attacker, SOC Analyst, External Alert "
     "Channels, Federated Nodes, and LLM API Provider.")
subsection_heading(doc, "3-2  Data Flow Overview")
body(doc,
     "Attacker traffic arrives at honeypots, events are normalized into AlertEvent "
     "records stored in alert_events, sessions are aggregated in cowrie_sessions, "
     "the behavior engine writes attacker_profiles, and all collections are served "
     "by the Flask API to dashboard clients via REST and WebSocket.")
subsection_heading(doc, "3-3  Context Diagram (DFD Level 0)")
placeholder_fig(doc, "A-1-3",
    "Insert the system context diagram showing NeuroTrap CADN as a single process "
    "with five external entities: Attacker, SOC Analyst, Alert Channels, Federated "
    "Nodes, and LLM API Provider. Label all data flows.",
    "Figure A-1-3. System Context Diagram (DFD Level 0)")
subsection_heading(doc, "3-4  DFD Level 0 - Main System Processes")
placeholder_fig(doc, "A-1-4",
    "Insert the Level 0 DFD decomposing the system into six major processes: "
    "(1) Honeypot Services, (2) Event Capture and Normalization, "
    "(3) Behavioral Analysis, (4) Deception and Response, "
    "(5) Intelligence Aggregation, (6) Dashboard and API. "
    "Show data flows between processes and data stores.",
    "Figure A-1-4. DFD Level 0 - Main System Processes")
subsection_heading(doc, "3-5  DFD Level 1 - Behavioral Analysis Decomposition")
placeholder_fig(doc, "A-1-5",
    "Insert the Level 1 DFD decomposing Behavioral Analysis: "
    "(3.1) Session Feature Extraction, (3.2) ML Intent Classification, "
    "(3.3) TTP Extraction, (3.4) Threat Score Computation, "
    "(3.5) Intent Reclassification, (3.6) Attacker Profile Write.",
    "Figure A-1-5. DFD Level 1 - Behavioral Analysis Process Decomposition")

# ── Section 4 ─────────────────────────────────────────────────────────────────
section_heading(doc, "4.  UML Diagrams")
subsection_heading(doc, "4-1  Use Case Diagram")
placeholder_fig(doc, "A-4-1",
    "Insert the Use Case Diagram with actors: Attacker, SOC Analyst, Admin. "
    "Attacker use cases: Connect to SSH/HTTP/Multi-Protocol Honeypot, Trigger "
    "Deception Environment, Receive Bait Content. "
    "SOC Analyst: View Dashboard, Threat Actors, MITRE Heatmap, CBEE, ADT, FHIM, SOC Reports. "
    "Admin: Login with MFA, Block IP, Generate Report, Recalculate Profiles.",
    "Figure A-4-1. System Use Case Diagram")
subsection_heading(doc, "4-2  Sequence Diagram")
placeholder_fig(doc, "A-4-2",
    "Insert the Sequence Diagram for attack detection. Lifelines: Attacker, "
    "Cowrie, Log Pipeline, Behavior Engine, MongoDB, Deception Engine, "
    "Response Engine, Flask API, Dashboard Client. Show 12-step sequence from "
    "SSH connection to WebSocket dashboard update.",
    "Figure A-4-2. Sequence Diagram - Attack Detection and Response Flow")
subsection_heading(doc, "4-3  Activity Diagram")
placeholder_fig(doc, "A-4-3",
    "Insert the Activity Diagram for threat response workflow showing decision "
    "branches for COWRIE_SKIP filter, severity assignment, session aggregation, "
    "threat score thresholds (>=90 block, >=70 isolate, >=40 slow, <40 log), "
    "bias injection threshold, and environment spawning.",
    "Figure A-4-3. Activity Diagram - Threat Response Workflow")

# ── Section 5 ─────────────────────────────────────────────────────────────────
section_heading(doc, "5.  Requirements Structuring")
subsection_heading(doc, "5-1  System Logic")
body(doc,
     "The core logic is a continuous event-driven processing loop. The system "
     "simultaneously listens for new TCP/UDP connections, tails Cowrie JSONL logs, "
     "processes the behavior engine queue, evaluates profiles against response "
     "thresholds, and serves REST and WebSocket requests. MongoDB is used as a "
     "shared state store using a decoupled producer-consumer pattern.")
subsection_heading(doc, "5-2  Structured English")
body(doc,
     "FOR EACH new session in cowrie_sessions WHERE processed = false: "
     "EXTRACT 13 features. CALL VotingClassifier. GET intent, tier, confidence. "
     "CALL TTPExtractor on commands. COMPUTE ttp_score. COMPUTE tier_bonus, "
     "persistence_bonus, volume_bonus. COMPUTE threat_score = (confidence x 40) "
     "+ ttp_score + tier_bonus + persistence_bonus + volume_bonus, capped at 100. "
     "CALL reclassify_intent on ALL accumulated commands. SAVE attacker_profile.")
subsection_heading(doc, "5-3  Decision Tables")
table_caption(doc, "Table A-5-1. Decision Table - Threat Response Matrix")
dt_rows = [
    ("threat_score >= 90",                 "Y", "N", "N", "N"),
    ("threat_score >= 70 and < 90",        "N", "Y", "N", "N"),
    ("threat_score >= 40 and < 70",        "N", "N", "Y", "N"),
    ("threat_score < 40",                  "N", "N", "N", "Y"),
    ("Execute iptables DROP",              "X", "-", "-", "-"),
    ("Capture PCAP (10,000 packets)",      "X", "-", "-", "-"),
    ("Execute iptables LOG + rate limit",  "-", "X", "-", "-"),
    ("Apply tc netem 200ms delay",         "-", "X", "-", "-"),
    ("Apply tc netem 500ms delay",         "-", "-", "X", "-"),
    ("Send SMTP/Slack/Telegram alert",     "X", "X", "-", "-"),
    ("Write to response_log",              "X", "X", "X", "X"),
]
add_table(doc, ["Condition/Action", "Rule 1", "Rule 2", "Rule 3", "Rule 4"], dt_rows)

subsection_heading(doc, "5-4  Decision Trees")
body(doc,
     "The intent reclassification tree evaluates all accumulated commands in "
     "priority order: (1) cryptomining - if xmrig/cryptonight/minerd/ethminer "
     "found; (2) malware_deployment - if wget/curl + chmod + execute pattern, "
     "or SCP upload + execute; (3) credential_harvesting - if failed_logins > 5 "
     "and unique usernames > 3; (4) bot_enrollment - if crontab/systemd/nohup "
     "persistence found; (5) lateral_movement - if SSH to new IP or subnet scan; "
     "(6) reconnaissance - default.")

# ── Section 6 ─────────────────────────────────────────────────────────────────
section_heading(doc, "6.  Conceptual Data Modeling")
subsection_heading(doc, "6-1  Entity Relationship Diagram")
placeholder_fig(doc, "A-6-1",
    "Insert the ERD showing entities: ALERT_EVENT, COWRIE_SESSION, "
    "ATTACKER_PROFILE, DECEPTION_ENVIRONMENT, CBEE_PROFILE, CBEE_INJECTION, "
    "GADCF_ASSET, ATTACKER_TWIN, RESPONSE_LOG, SOC_REPORT, FHIM_ROUND, "
    "HONEYPOT_SESSION. Show key attributes and relationships (all related to "
    "ATTACKER_PROFILE via src_ip foreign key).",
    "Figure A-6-1. Entity Relationship Diagram")
subsection_heading(doc, "6-2  Physical Data Model")
body(doc,
     "The physical model is implemented in MongoDB 6.0. Indexes created by "
     "setup_db_indexes.py: alert_events (src_ip, timestamp desc, attack_type, "
     "compound src_ip+timestamp); cowrie_sessions (src_ip, processed); "
     "attacker_profiles (threat_score desc, src_ip); response_log (timestamp desc, action).")
subsection_heading(doc, "6-3  Logical Data Model")
body(doc,
     "The central entity is attacker_profile identified by src_ip. All time-series "
     "collections (alert_events, cowrie_sessions, response_log, cbee_injections) are "
     "queried sorted by timestamp descending. The behavior engine uses an upsert "
     "pattern: create new profile or increment session_count and recalculate score.")
subsection_heading(doc, "6-4  Class Diagram")
placeholder_fig(doc, "A-6-4",
    "Insert the UML Class Diagram showing key classes: AlertEvent, "
    "AttackerProfile, ProfileStore, SessionFeatureExtractor, AttackerClassifier, "
    "TTPExtractor, DeceptionEngine, CredentialGenerator, CBEEEngine, BiasProfile, "
    "GADCFEngine, DigitalTwin, TacticPredictor, KillChainMapper, FederatedNode, "
    "FedAvgServer, ResponseEngine, SOCAnalyst. Show attributes, methods, "
    "and associations.",
    "Figure A-6-4. System Class Diagram")
subsection_heading(doc, "6-5  Business Rules")
table_caption(doc, "Table A-6-5. Business Rules")
br_rows = [
    ("BR-01","Port 22 exclusively used by Cowrie; management SSH on port 50402.","UFW + Docker port mapping"),
    ("BR-02","No container runs in privileged mode.","Docker Compose config"),
    ("BR-03","NET_ADMIN/NET_RAW only for packet-monitor container.","Docker Compose capabilities"),
    ("BR-04","MongoDB port 27017 never exposed to internet.","Docker internal network isolation"),
    ("BR-05","All secrets in .env files not tracked by git.","gitignore rules"),
    ("BR-06","Default credentials must be changed before production deployment.","Security documentation"),
    ("BR-07","Deception environments auto-terminated after 60 minutes.","DeceptionEngine timeout"),
    ("BR-08","Maximum 20 simultaneous active deception environments.","DeceptionEngine counter"),
    ("BR-09","CBEE bait injection limited to 3 per attacker IP.","MAX_INJECTIONS_PER_IP constant"),
    ("BR-10","Intent reclassification evaluates all accumulated commands.","reclassify_intent() full scan"),
    ("BR-11","Analyst role is read-only; admin role has full write access.","Flask JWT role check"),
    ("BR-12","All response actions written to response_log regardless of success.","Mock mode fallback"),
]
add_table(doc, ["ID", "Rule", "Enforcement"], br_rows)


# ══════════════════════════════════════════════════════════════════════════════
# ═══════════════════════  DESIGN PHASE  ══════════════════════════════════════
# ══════════════════════════════════════════════════════════════════════════════
add_page_break(doc)
chapter_heading(doc, "DESIGN PHASE")

# ── Section 1 ─────────────────────────────────────────────────────────────────
section_heading(doc, "1.  System Design")
subsection_heading(doc, "1-1  Overview")
body(doc,
     "The system design translates requirements and logical models into concrete "
     "architectural specifications. The design follows three principles: separation "
     "of concerns (each container handles one domain), defense in depth (multiple "
     "security layers), and graceful degradation (reduced-function mode when "
     "optional dependencies are unavailable).")
subsection_heading(doc, "1-2  Component Design Summary")
table_caption(doc, "Table D-1-1. System Component Design Summary")
comp_rows = [
    ("neurotrap-cowrie",    "cowrie/cowrie:latest",      "honeypot-net, elk-net",                           "High-interaction SSH/Telnet emulation"),
    ("neurotrap-opencanary","custom Dockerfile",         "honeypot-net, elk-net",                           "Low-interaction multi-protocol (8 services)"),
    ("neurotrap-galah",     "0x4d31/galah:latest",       "honeypot-net, elk-net",                           "LLM-powered dynamic HTTP responses"),
    ("neurotrap-monitor",   "custom Dockerfile",         "host network",                                    "Scapy packet capture + JSONL tail"),
    ("neurotrap-behavior",  "custom Dockerfile",         "elk-net, management-net",                         "ML classifier + threat scorer loop"),
    ("neurotrap-deception", "custom Dockerfile",         "honeypot-net, elk-net, management-net",           "Template-driven environment spawner"),
    ("neurotrap-mongodb",   "mongo:6.0",                 "elk-net, management-net, monitor-bridge (static 172.25.0.10)", "Central document store"),
    ("neurotrap-api",       "custom Dockerfile",         "management-net, monitor-bridge",                  "JWT REST API + WebSocket + dashboard"),
    ("neurotrap-nginx",     "nginx:alpine",              "honeypot-net, management-net",                    "TLS termination + reverse proxy"),
]
add_table(doc, ["Container", "Image", "Networks", "Role"], comp_rows)

subsection_heading(doc, "1-3  Network Architecture Design")
body(doc,
     "Four Docker virtual networks enforce strict traffic segmentation. "
     "honeypot-net (172.20.0.0/24) carries attacker traffic. "
     "elk-net is an internal bridge with no gateway for log transport. "
     "management-net is an internal bridge for API/database communication. "
     "monitor-bridge (172.25.0.0/24, non-internal) connects the host-mode "
     "packet-monitor container to MongoDB at static IP 172.25.0.10.")
subsection_heading(doc, "1-4  Behavioral Analysis Design")
body(doc,
     "The ML pipeline uses a VotingClassifier (RandomForest + SVC, soft voting) "
     "on a 13-dimensional session-level feature vector: total_commands, "
     "unique_commands, dangerous_count, recon_count, download_attempts, "
     "file_access, session_duration, login_attempts, failed_logins, "
     "has_persistence, has_lateral, dangerous_ratio, and recon_ratio.")
body(doc,
     "Threat score formula: score = (ML_confidence x 40) + ttp_score + tier_bonus "
     "+ persistence_bonus + volume_bonus, capped at 100. Tier bonuses: "
     "beginner=0, automated_bot=15, advanced_human=30. Persistence bonus: "
     "1 session=5, 2=18, 3-4=22, 5-9=28, 10-19=40, 20-49=50, 50-99=60, 100+=65. "
     "Volume bonus: min(total_commands // 5, 15).")
subsection_heading(doc, "1-5  Deception Environment Design")
body(doc,
     "Three environment templates are defined: "
     "beginner (simple-linux, Ubuntu 20.04 workstation), "
     "automated_bot (baited-server, CentOS 7 with database credentials and "
     "application configuration files), and "
     "advanced_human (advanced-corporate, Debian 11 with git repositories, "
     "AWS credentials, .env files, and fake corporate email threads). "
     "CredentialGenerator uses Faker seeded by attacker IP for deterministic, "
     "internally consistent fake credentials.")

# ── Section 2 ─────────────────────────────────────────────────────────────────
section_heading(doc, "2.  Menu Design")
subsection_heading(doc, "2-1  Dashboard Navigation Architecture")
body(doc,
     "The dashboard is organized into two navigation groups. Operations group: "
     "Dashboard, Threat Actors, Live Events, Honeypots, Response Log. "
     "Intelligence group: Threat Intel, Geo Map, MITRE ATT&CK, Behavior Analysis, "
     "CBEE, GADCF, FHIM, ADT, AI Analyst.")
placeholder_fig(doc, "D-2-1",
    "Insert a screenshot or wireframe of the dashboard main navigation menu, "
    "showing the left sidebar with Operations and Intelligence groups and all "
    "listed sections with their icons.",
    "Figure D-2-1. Dashboard Main Navigation Menu")

# ── Section 3 ─────────────────────────────────────────────────────────────────
section_heading(doc, "3.  Forms Design")
subsection_heading(doc, "3-1  Login Form")
body(doc,
     "The login form submits POST /api/auth/login with JSON body containing "
     "username, password, and optionally otp. On success the server returns a "
     "JWT token stored in localStorage. When MFA_ENABLED=1, the OTP field is "
     "required and displayed after the form calls /api/auth/mfa/status on load.")
placeholder_fig(doc, "D-3-1",
    "Insert screenshot of the login form showing: username field, password field "
    "(masked), OTP field (shown when MFA enabled), and Login button, centered "
    "on the page with the NeuroTrap CADN logo above.",
    "Figure D-3-1. Login Form Design")
subsection_heading(doc, "3-2  Attacker Profile Detail Modal")
placeholder_fig(doc, "D-3-2",
    "Insert screenshot of the attacker profile modal showing: IP address, "
    "country flag, threat score gauge, intent badge, tier badge, session count, "
    "last seen timestamp, top commands table, MITRE techniques list, "
    "5-dimension bias bar chart, and Block IP button (admin only).",
    "Figure D-3-2. Attacker Profile Detail Modal")
subsection_heading(doc, "3-3  CBEE Bias Scorer Form")
placeholder_fig(doc, "D-3-3",
    "Insert screenshot of the CBEE ad-hoc scoring form (admin only, "
    "POST /api/cbee/score) showing: IP address input, session commands text area, "
    "Score Biases button, and returned bias dimension scores as radar chart.",
    "Figure D-3-3. CBEE Cognitive Bias Scorer Form")

# ── Section 4 ─────────────────────────────────────────────────────────────────
section_heading(doc, "4.  Reports Design")
subsection_heading(doc, "4-1  SOC Incident Report Format")
body(doc,
     "Incident reports are generated by the SOCAnalyst module using LLM mode "
     "(when ANTHROPIC_API_KEY or GROQ_API_KEY is configured) or heuristic "
     "template filling (fallback). Reports are stored in soc_reports and "
     "accessible via GET /api/soc/reports.")
placeholder_fig(doc, "D-4-1",
    "Insert screenshot of a generated SOC incident report showing sections: "
    "Executive Summary, Incident Timeline table (timestamps + event types), "
    "MITRE ATT&CK Coverage Table (technique IDs, names, tactic, confidence), "
    "Risk Assessment (threat score, tier, intent), and Recommended Response list.",
    "Figure D-4-1. SOC Incident Report Template")


# ══════════════════════════════════════════════════════════════════════════════
# ═══════════════════  IMPLEMENTATION PHASE  ══════════════════════════════════
# ══════════════════════════════════════════════════════════════════════════════
add_page_break(doc)
chapter_heading(doc, "IMPLEMENTATION PHASE")

# ── Section 1 ─────────────────────────────────────────────────────────────────
section_heading(doc, "1.  System Deployment")
subsection_heading(doc, "1-1  Deployment Environment")
body(doc,
     "NeuroTrap CADN is deployed on a VPS with Ubuntu 24.04 LTS, 6 vCPUs, "
     "11 GB RAM, and 193 GB SSD at public IP 13.140.144.118. Real SSH management "
     "is on port 50402 to avoid conflict with the Cowrie honeypot on port 22.")
subsection_heading(doc, "1-2  Container Deployment")
table_caption(doc, "Table I-1-1. Container Deployment Summary")
deploy_rows = [
    ("neurotrap-nginx",     "nginx:alpine",             "443:443, 8080:80",         "unless-stopped"),
    ("neurotrap-api",       "custom Dockerfile.api",    "5000:5000",               "unless-stopped"),
    ("neurotrap-cowrie",    "cowrie/cowrie:latest",     "22:2222, 23:2223",        "unless-stopped"),
    ("neurotrap-opencanary","custom Dockerfile",        "21, 80, 445, 3306, 1433, 161/udp, 5900, 3389", "unless-stopped"),
    ("neurotrap-galah",     "0x4d31/galah:latest",      "8088:8080",               "unless-stopped"),
    ("neurotrap-mongodb",   "mongo:6.0",                "internal only",           "unless-stopped"),
    ("neurotrap-monitor",   "custom Dockerfile",        "host network",            "unless-stopped"),
    ("neurotrap-behavior",  "custom Dockerfile",        "internal only",           "unless-stopped"),
    ("neurotrap-deception", "custom Dockerfile",        "internal only",           "unless-stopped"),
]
add_table(doc, ["Container", "Image", "Ports", "Restart"], deploy_rows)

subsection_heading(doc, "1-3  Key Applied Fixes")
fixes = [
    ("Fix 1 - Monitor-MongoDB Connectivity",
     "The packet-monitor uses host network mode for raw packet capture. "
     "A dedicated monitor-bridge network (172.25.0.0/24) assigns MongoDB "
     "static IP 172.25.0.10 accessible from the host network."),
    ("Fix 2 - Cowrie Log Volume Path",
     "Cowrie logs land at /cowrie/cowrie-git/var/log/cowrie/. The volume "
     "was incorrectly mounted at /cowrie/var/log/cowrie. Fixed to correct path."),
    ("Fix 8 - Deception Environment Threshold",
     "Original threshold was threat_score >= 30. Real attacker scores clustered "
     "at 15-25. Threshold lowered to >= 10."),
    ("Fix 11 - Threat Score Formula Recalibration",
     "Original formula produced scores clustered in LOW band. Formula rewritten "
     "with tiered persistence bonus (5-65 points) and volume bonus (up to 15 points)."),
    ("Fix 12 - Intent Reclassification",
     "ML fallback defaulted to reconnaissance. reclassify_intent() added to "
     "re-examine all accumulated session commands with priority-ordered rules."),
    ("Fix 14 - CBEE Background Thread",
     "CBEEEngine was created but .start() was never called, so the background "
     "scoring thread never ran. Fixed by calling engine.start() after creation."),
]
for title, desc in fixes:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run1 = p.add_run(title + ": ")
    run1.bold = True; run1.font.name='Times New Roman'; run1.font.size=Pt(12)
    run2 = p.add_run(desc)
    run2.font.name='Times New Roman'; run2.font.size=Pt(12)

subsection_heading(doc, "1-4  Environment Configuration")
body(doc,
     "Required .env variables: MONGO_USER, MONGO_PASS, MONGO_URI, SECRET_KEY, "
     "JWT_SECRET, ADMIN_USER, ADMIN_PASS, ANALYST_USER, ANALYST_PASS. "
     "Optional variables: MFA_ENABLED, MFA_SECRET, ANTHROPIC_API_KEY, "
     "GROQ_API_KEY, SOC_ANALYST_MODEL, MONITOR_INTERFACE, SLACK_WEBHOOK_URL, "
     "TELEGRAM_TOKEN, TELEGRAM_CHAT_ID, SMTP_HOST.")

subsection_heading(doc, "1-5  Deployment Procedure")
steps = [
    "Clone the repository and navigate to the project directory.",
    "Create .env file with all required variables. Generate secure random values "
    "with: python -c \"import secrets; print(secrets.token_hex(32))\"",
    "Generate self-signed SSL certificate: bash scripts/generate_ssl_cert.sh",
    "Configure UFW to allow honeypot ports (22, 23, 21, 80, 443, 445, 1433, 3306, "
    "3389, 5900, 8088, 161/udp) and management SSH port 50402.",
    "Build and start all containers: docker compose up -d --build",
    "Initialize MongoDB indexes: docker compose exec api python -m src.database.setup_db_indexes",
    "Access dashboard at https://[SERVER_IP] using admin credentials.",
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(step)
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(4)

# ── Section 2 ─────────────────────────────────────────────────────────────────
section_heading(doc, "2.  System Snapshots")
body(doc,
     "The following figures show the running NeuroTrap CADN system dashboard. "
     "All figures require actual screenshots from the deployed production system.")

snap_figs = [
    ("I-2-1",  "Dashboard Main View - KPI Overview",
     "Insert screenshot showing four KPI cards (Total Events, Active Sessions, "
     "IPs Blocked, Active Environments), Attack Type Distribution bar chart, "
     "Threat Level Distribution pie chart, and Live Events feed."),
    ("I-2-2",  "Live Events Feed Panel",
     "Insert screenshot of Live Events feed showing columns: timestamp, source IP, "
     "destination port, attack type badge (color-coded), honeypot source, severity, "
     "and events-per-minute counter."),
    ("I-2-3",  "Threat Actors Table",
     "Insert screenshot of Threat Actors table ranked by threat score with columns: "
     "IP, country flag, threat score bar, intent, tier badge, session count, "
     "last seen, View Profile button."),
    ("I-2-4",  "Honeypots Status Panel",
     "Insert screenshot showing three sub-sections: Live Sensors (hit counts + status), "
     "Recent Attacker Sessions table, and Active Deception Environments list."),
    ("I-2-5",  "CBEE Cognitive Bias Panel",
     "Insert screenshot showing cognitive bias profiles with radar charts of five "
     "dimensions per attacker and dominant bias type highlighted."),
    ("I-2-6",  "GADCF Asset Generation Panel",
     "Insert screenshot showing generated deception assets with type, target IP, "
     "industry category, and generation timestamp."),
    ("I-2-7",  "Attacker Digital Twin Detail Panel",
     "Insert screenshot showing kill chain stage indicator, predicted next tactics "
     "(top 3), fidelity score, automation score, and MITRE tactic distribution."),
    ("I-2-8",  "FHIM Federated Learning Panel",
     "Insert screenshot showing four demo nodes (cairo-uni-01, acme-financial-01, "
     "fraunhofer-fkie-01, saudi-telecom-01) with status and aggregation round history."),
    ("I-2-9",  "AI SOC Analyst Panel",
     "Insert screenshot showing triage queue, incident report for one attacker, "
     "and Q&A chat interface with example question and AI-generated response."),
    ("I-2-10", "MITRE ATT&CK Technique Frequency Heatmap",
     "Insert screenshot of the MITRE ATT&CK heatmap with technique IDs on horizontal "
     "axis, tactic categories on vertical axis, and cell color intensity showing frequency."),
    ("I-2-11", "Attacker Geolocation Map",
     "Insert screenshot of the Leaflet.js interactive world map with clustered "
     "attacker IP geolocation markers."),
    ("I-2-12", "System Architecture Deployment Diagram",
     "Insert UML Deployment Diagram showing the VPS node, nine container nodes, "
     "four Docker networks, attacker client, and SOC analyst browser."),
]
for fig_id, fig_title, description in snap_figs:
    placeholder_fig(doc, fig_id, description, f"Figure {fig_id}. {fig_title}")

# ── Section 3 ─────────────────────────────────────────────────────────────────
section_heading(doc, "3.  System Maintenance and Support")
subsection_heading(doc, "3-1  Routine Maintenance Schedule")
for period, tasks in [
    ("Daily", "Review SOC Analyst shift summary (GET /api/soc/summary). "
              "Check triage queue for CRITICAL items. "
              "Monitor container health with docker ps."),
    ("Weekly", "Review response log for patterns (GET /api/response/log). "
               "Check disk usage with docker system df. "
               "Recalculate attacker profiles (POST /api/profiles/recalculate)."),
    ("Monthly", "Rotate JWT_SECRET and ADMIN_PASS. "
                "Review and update UFW rules. "
                "Check Docker image updates and rebuild containers."),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run1 = p.add_run(period + ": ")
    run1.bold = True; run1.font.name='Times New Roman'; run1.font.size=Pt(12)
    run2 = p.add_run(tasks)
    run2.font.name='Times New Roman'; run2.font.size=Pt(12)

subsection_heading(doc, "3-2  Common Troubleshooting")
issues = [
    ("502 Bad Gateway after API restart",
     "Reload nginx: docker compose exec nginx nginx -s reload"),
    ("Behavior engine not processing sessions",
     "Verify cowrie-logs volume is correctly mounted in Cowrie and monitor containers."),
    ("All threat scores LOW",
     "Recalculate profiles: curl -X POST http://localhost:5000/api/profiles/recalculate"),
    ("CBEE profiles loading forever",
     "Check for boolean evaluation of PyMongo database object; ensure db is not None check."),
]
for issue, resolution in issues:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run1 = p.add_run(issue + ": ")
    run1.bold = True; run1.font.name='Times New Roman'; run1.font.size=Pt(12)
    run2 = p.add_run(resolution)
    run2.font.name='Times New Roman'; run2.font.size=Pt(12)

# ── Section 4 ─────────────────────────────────────────────────────────────────
section_heading(doc, "4.  Non-Functional Requirements Implementation")
table_caption(doc, "Table I-4-1. Non-Functional Requirements Verification")
nfr_ver_rows = [
    ("NFR-01","Event-to-dashboard latency","< 5s",   "[PLACEHOLDER]","[PLACEHOLDER]"),
    ("NFR-02","ML classifier macro F1 score","> 0.85","[PLACEHOLDER]","[PLACEHOLDER]"),
    ("NFR-03","Environment spawn time","< 30s",      "[PLACEHOLDER]","[PLACEHOLDER]"),
    ("NFR-04","Response action time","< 10s",        "[PLACEHOLDER]","[PLACEHOLDER]"),
    ("NFR-05","Lynis hardening score","> 70",        "[PLACEHOLDER]","[PLACEHOLDER]"),
    ("NFR-06","Max simultaneous environments","20",  "Enforced by code","Implemented"),
    ("NFR-07","Environment lifetime","60 min",       "Enforced by code","Implemented"),
    ("NFR-08","MongoDB availability","99%+",         "restart: unless-stopped","Implemented"),
    ("NFR-09","API response cache","30 s",           "@cached decorator","Implemented"),
    ("NFR-10","JWT token expiry","1 hour",           "flask-jwt-extended","Implemented"),
    ("NFR-11","FHIM DP epsilon","1.0",               "FederatedNode constant","Implemented"),
    ("NFR-12","FHIM DP delta","1e-5",                "FederatedNode constant","Implemented"),
]
add_table(doc, ["ID","Requirement","Target","Actual","Status"], nfr_ver_rows)

subsection_heading(doc, "4-1  Performance Optimization")
body(doc,
     "API response caching uses an in-memory dictionary keyed by endpoint path "
     "and query string, storing serialized JSON for 30 seconds. WebSocket-based "
     "live event delivery broadcasts new_event messages to all subscribed clients "
     "immediately on AlertEvent creation, achieving sub-second delivery.")
subsection_heading(doc, "4-2  Security Implementation")
body(doc,
     "Four security areas are implemented: (1) Network Isolation via four Docker "
     "networks with no internet-facing internal networks. (2) Authentication via "
     "JWT tokens with one-hour expiry and optional TOTP MFA for admin. "
     "(3) Container Hardening: no privileged mode, NET_ADMIN/NET_RAW only for "
     "packet-monitor, Cowrie as uid 999, read-only config mounts. "
     "(4) Host Hardening: SSH on port 50402, UFW rules, target Lynis score > 70.")


# ══════════════════════════════════════════════════════════════════════════════
# CONCLUSION AND FUTURE WORK
# ══════════════════════════════════════════════════════════════════════════════
add_page_break(doc)
chapter_heading(doc, "CONCLUSION AND FUTURE WORK")

section_heading(doc, "Conclusion")
body(doc,
     "NeuroTrap CADN demonstrates that it is possible to build a comprehensive, "
     "production-grade active defense platform from open-source components at "
     "zero software licensing cost. The system successfully integrates eleven "
     "network honeypot services, a machine learning behavioral analysis pipeline, "
     "cognitive psychology-driven deception mechanisms, and federated privacy-"
     "preserving intelligence sharing into a unified platform accessible through "
     "a real-time operational dashboard.")
body(doc,
     "The ten-layer processing architecture provides clean separation between event "
     "capture, behavioral analysis, deception generation, and response action, "
     "enabling each layer to be enhanced or replaced independently. The Cognitive "
     "Bias Exploitation Engine introduces a novel dimension to honeypot deception "
     "by quantifying attacker psychological state across five bias dimensions and "
     "using that profile to select contextually relevant deceptive content.")
body(doc,
     "The Federated Honeypot Intelligence Mesh addresses the challenge of cross-"
     "organizational threat intelligence sharing by applying federated averaging "
     "with Gaussian differential privacy noise, providing a practical foundation "
     "for real multi-organizational deployment.")

section_heading(doc, "Results")
body(doc,
     "[PLACEHOLDER: Insert actual measured results after running the formal "
     "evaluation test suite. Replace the following with measured values:]")
results = [
    "ML classifier macro F1 score: [PLACEHOLDER - target > 0.85]",
    "Event-to-dashboard latency: [PLACEHOLDER - target < 5 seconds]",
    "Deception environment spawn time: [PLACEHOLDER - target < 30 seconds]",
    "Response action execution time: [PLACEHOLDER - target < 10 seconds]",
    "Lynis hardening score: [PLACEHOLDER - target > 70]",
    "Total honeypot events captured during test period: [PLACEHOLDER]",
    "Unique attacker IPs profiled: [PLACEHOLDER]",
    "Deception environments spawned: [PLACEHOLDER]",
    "Response actions executed: [PLACEHOLDER]",
]
for item in results:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(3)

section_heading(doc, "Recommendations for Future Work")
future_work = [
    ("ASHRTA Module Implementation",
     "The planned ASHRTA (Advanced Stealth Honeypot Reconnaissance and Tracking "
     "Adaptation) module would add anti-fingerprinting capabilities by dynamically "
     "modifying service banners, response timing, and filesystem artifacts."),
    ("Real Federated Learning Deployment",
     "Establish real federated connections between multiple deployed NeuroTrap "
     "instances with TLS-secured inter-node communication and authenticated delta "
     "submission."),
    ("Windows Honeypot Integration",
     "Add Windows-specific honeypot services (Active Directory emulation, Windows "
     "RDP with session capture, Office macro sandbox) to extend coverage."),
    ("Load Testing and Performance Profiling",
     "Conduct structured load testing measuring events-per-second capacity, MongoDB "
     "write throughput, and WebSocket latency under concurrent connections."),
    ("SIEM Integration",
     "Add support for forwarding normalized events to external SIEM platforms via "
     "CEF, LEEF, or STIX/TAXII formats."),
    ("Real Training Data",
     "Collect real attacker session data from the production deployment to retrain "
     "the ML classifier for improved accuracy."),
    ("Mobile SOC Application",
     "Develop a native mobile application for the SOC analyst interface enabling "
     "on-call notification response without laptop access."),
]
for i, (title, desc) in enumerate(future_work, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run1 = p.add_run(f"{i}. {title}: ")
    run1.bold = True; run1.font.name='Times New Roman'; run1.font.size=Pt(12)
    run2 = p.add_run(desc)
    run2.font.name='Times New Roman'; run2.font.size=Pt(12)


# ══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
add_page_break(doc)
chapter_heading(doc, "REFERENCES")

references = [
    "[1] L. Spitzner, Honeypots: Tracking Hackers. Boston: Addison-Wesley, 2002.",
    "[2] MITRE Corporation, (2024). MITRE ATT&CK: Adversarial Tactics, Techniques, "
    "and Common Knowledge. Available at: https://attack.mitre.org. "
    "Last visit date: 01/06/2026.",
    "[3] M. Oosterhof, (2024). Cowrie SSH/Telnet Honeypot. Available at: "
    "https://github.com/cowrie/cowrie. Last visit date: 01/06/2026.",
    "[4] Thinkst Applied Research, (2024). OpenCanary: Multi-Protocol Network Honeypot. "
    "Available at: https://github.com/thinkst/opencanary. Last visit date: 01/06/2026.",
    "[5] 0x4D31, (2024). Galah: An LLM-Powered Web Honeypot. Available at: "
    "https://github.com/0x4D31/galah. Last visit date: 01/06/2026.",
    "[6] P. Biondi et al., (2024). Scapy: Packet Manipulation Library for Python. "
    "Available at: https://scapy.net. Last visit date: 01/06/2026.",
    "[7] F. Pedregosa et al., \"Scikit-learn: Machine Learning in Python\", "
    "Journal of Machine Learning Research, Vol. 12, 2011, pp. 2825-2830.",
    "[8] N. Reimers and I. Gurevych, \"Sentence-BERT: Sentence Embeddings using "
    "Siamese BERT-Networks\", Proceedings of EMNLP, 2019.",
    "[9] H. B. McMahan et al., \"Communication-Efficient Learning of Deep Networks "
    "from Decentralized Data\", Proceedings of AISTATS, 2017.",
    "[10] C. Dwork and A. Roth, \"The Algorithmic Foundations of Differential "
    "Privacy\", Foundations and Trends in Theoretical Computer Science, "
    "Vol. 9, No. 3-4, 2014, pp. 211-407.",
    "[11] E. M. Hutchins, M. J. Cloppert, and R. M. Amin, \"Intelligence-Driven "
    "Computer Network Defense Informed by Analysis of Adversary Campaigns and "
    "Intrusion Kill Chains\", Proceedings of the 6th ICIS, 2011.",
    "[12] A. Ronacher, (2024). Flask: A Lightweight WSGI Web Application Framework. "
    "Available at: https://flask.palletsprojects.com. Last visit date: 01/06/2026.",
    "[13] MongoDB, Inc., (2024). MongoDB 6.0 Documentation. Available at: "
    "https://www.mongodb.com/docs. Last visit date: 01/06/2026.",
    "[14] Docker, Inc., (2024). Docker Engine Documentation. Available at: "
    "https://docs.docker.com. Last visit date: 01/06/2026.",
    "[15] MongoDB, Inc., (2024). PyMongo: Python Driver for MongoDB. Available at: "
    "https://pymongo.readthedocs.io. Last visit date: 01/06/2026.",
    "[16] Viro Solutions, (2024). Flask-JWT-Extended: JWT Manager for Flask. "
    "Available at: https://flask-jwt-extended.readthedocs.io. Last visit date: 01/06/2026.",
    "[17] Anthropic, (2024). Claude API Documentation. Available at: "
    "https://docs.anthropic.com. Last visit date: 01/06/2026.",
    "[18] B. Faraglia, (2024). Faker: A Python Package to Generate Fake Data. "
    "Available at: https://faker.readthedocs.io. Last visit date: 01/06/2026.",
    "[19] H. Totp, (2024). PyOTP: Python One-Time Password Library. Available at: "
    "https://pyauth.github.io/pyotp. Last visit date: 01/06/2026.",
    "[20] Ollama, (2024). Ollama: Get Up and Running with Large Language Models "
    "Locally. Available at: https://ollama.com. Last visit date: 01/06/2026.",
]
for ref in references:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after  = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.left_indent  = Cm(0.6)
    pf.first_line_indent = Cm(-0.6)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)


# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX A
# ══════════════════════════════════════════════════════════════════════════════
add_page_break(doc)
chapter_heading(doc, "APPENDIX A - SOURCE CODE EXCERPTS")
para(doc, "Page A1", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.RIGHT,
     space_before=0, space_after=12)

section_heading(doc, "A-1  AlertEvent Schema (src/detection/alert_schema.py)")
code_text = (
    "from dataclasses import dataclass, field\n"
    "from enum import Enum\n"
    "from datetime import datetime\n"
    "import uuid\n\n"
    "class ATTACK_TYPES(str, Enum):\n"
    "    brute_force       = 'brute_force'\n"
    "    command_injection = 'command_injection'\n"
    "    tool_fingerprint  = 'tool_fingerprint'\n"
    "    protocol_anomaly  = 'protocol_anomaly'\n"
    "    lateral_movement  = 'lateral_movement'\n"
    "    malware_upload    = 'malware_upload'\n"
    "    port_scan         = 'port_scan'\n"
    "    unknown           = 'unknown'\n\n"
    "_COWRIE_SKIP = frozenset({\n"
    "    'cowrie.client.kex', 'cowrie.client.var',\n"
    "    'cowrie.client.fingerprint', 'cowrie.connection.failed',\n"
    "    'cowrie.session.connect',\n"
    "})\n\n"
    "@dataclass\n"
    "class AlertEvent:\n"
    "    src_ip: str\n"
    "    dst_port: int\n"
    "    attack_type: str\n"
    "    honeypot_source: str\n"
    "    severity: str\n"
    "    src_port: int = 0\n"
    "    session_id: str = ''\n"
    "    username: str = ''\n"
    "    password: str = ''\n"
    "    command: str = ''\n"
    "    raw_payload: str = ''\n"
    "    timestamp: str = field(\n"
    "        default_factory=lambda: datetime.utcnow().isoformat())\n"
    "    event_id: str = field(\n"
    "        default_factory=lambda: str(uuid.uuid4()))"
)
p = doc.add_paragraph()
run = p.add_run(code_text)
run.font.name = 'Courier New'
run.font.size = Pt(10)
p.paragraph_format.space_after = Pt(12)

para(doc, "Page A2", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.RIGHT,
     space_before=12, space_after=12)
section_heading(doc, "A-2  Threat Score Formula (src/behavior/attacker_profile.py)")
code_text2 = (
    "def _compute_threat_score(self) -> float:\n"
    "    PERSISTENCE_BONUS = {\n"
    "        1:5, 2:18, 3:22, 4:22, 5:28, 9:28,\n"
    "        10:40, 19:40, 20:50, 49:50, 50:60, 99:60, 100:65,\n"
    "    }\n"
    "    def _persistence(n):\n"
    "        for t in sorted(PERSISTENCE_BONUS, reverse=True):\n"
    "            if n >= t: return PERSISTENCE_BONUS[t]\n"
    "        return 5\n\n"
    "    TACTIC_WEIGHTS = {\n"
    "        'Impact':40, 'PrivilegeEscalation':35, 'CredentialAccess':30,\n"
    "        'LateralMovement':25, 'Persistence':20, 'CommandAndControl':15,\n"
    "        'DefenseEvasion':10, 'Discovery':5,\n"
    "    }\n"
    "    TIER_BONUS = {'beginner':0, 'automated_bot':15, 'advanced_human':30}\n\n"
    "    base          = self.ml_confidence * 40\n"
    "    ttp_score     = min(sum(TACTIC_WEIGHTS.get(t.tactic,5)*t.confidence\n"
    "                           for t in self.ttps), 40)\n"
    "    tier_bonus    = TIER_BONUS.get(self.tier, 0)\n"
    "    persist_bonus = _persistence(self.session_count)\n"
    "    vol_bonus     = min(self.total_commands // 5, 15)\n"
    "    return min(base+ttp_score+tier_bonus+persist_bonus+vol_bonus, 100)"
)
p = doc.add_paragraph()
run = p.add_run(code_text2)
run.font.name = 'Courier New'
run.font.size = Pt(10)
p.paragraph_format.space_after = Pt(0)


# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX B
# ══════════════════════════════════════════════════════════════════════════════
add_page_break(doc)
chapter_heading(doc, "APPENDIX B - API REFERENCE")
para(doc, "Page B1", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.RIGHT,
     space_before=0, space_after=12)

section_heading(doc, "B-1  Authentication Endpoints")
table_caption(doc, "Table B-1-1. Authentication Endpoints")
auth_rows = [
    ("POST", "/api/auth/login",         "None",  "Authenticate and receive JWT token"),
    ("GET",  "/api/auth/mfa/status",    "None",  "Check MFA enabled/configured flags"),
    ("GET",  "/api/auth/otp/setup",     "Admin", "Generate new TOTP secret and QR URI"),
    ("POST", "/api/auth/otp/verify",    "None",  "Verify a TOTP code pre-login"),
]
add_table(doc, ["Method","Endpoint","Auth","Description"], auth_rows)

section_heading(doc, "B-2  Core Data Endpoints")
table_caption(doc, "Table B-2-1. Core Data Endpoints")
core_rows = [
    ("GET",  "/api/events",              "None",  "Alert events with filters"),
    ("GET",  "/api/events/stats",        "None",  "KPI counts and attack type distribution"),
    ("GET",  "/api/attackers",           "None",  "Top attacker profiles by threat score"),
    ("GET",  "/api/attackers/<src_ip>",  "None",  "Single attacker profile detail"),
    ("POST", "/api/profiles/recalculate","Admin", "Recalculate all profiles"),
    ("GET",  "/api/honeypots",           "None",  "Sensor hit counts and recent sessions"),
    ("GET",  "/api/environments",        "None",  "All deception environments"),
    ("GET",  "/api/intel",               "None",  "IOC list, top countries, campaigns"),
]
add_table(doc, ["Method","Endpoint","Auth","Description"], core_rows)

section_heading(doc, "B-3  Intelligence and Response Endpoints")
table_caption(doc, "Table B-3-1. Intelligence and Response Endpoints")
intel_rows = [
    ("GET",  "/api/cbee/profiles",       "None",  "Cognitive bias profiles"),
    ("GET",  "/api/cbee/injections",     "None",  "Bait injection log"),
    ("POST", "/api/cbee/score",          "Admin", "Ad-hoc bias scoring"),
    ("GET",  "/api/gadcf/assets",        "None",  "Generated deception assets"),
    ("POST", "/api/gadcf/generate",      "Admin", "Trigger content generation"),
    ("GET",  "/api/fhim/nodes",          "None",  "Federated node status"),
    ("GET",  "/api/twin/list",           "None",  "All attacker digital twins"),
    ("GET",  "/api/soc/summary",         "None",  "AI SOC shift summary"),
    ("GET",  "/api/soc/triage",          "None",  "Ranked incident queue"),
    ("POST", "/api/soc/chat",            "Admin", "Analyst Q&A"),
    ("POST", "/api/response/block",      "Admin", "Manual IP block"),
    ("GET",  "/api/response/log",        "None",  "Response action history"),
]
add_table(doc, ["Method","Endpoint","Auth","Description"], intel_rows)


# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX C
# ══════════════════════════════════════════════════════════════════════════════
add_page_break(doc)
chapter_heading(doc, "APPENDIX C - INSTALLATION GUIDE")
para(doc, "Page C1", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.RIGHT,
     space_before=0, space_after=12)

section_heading(doc, "C-1  Prerequisites")
body(doc,
     "Ubuntu 22.04 LTS or 24.04 LTS, minimum 4-core CPU / 8 GB RAM / 50 GB SSD, "
     "public IP address, Docker Engine 24.0+, Docker Compose 2.x, Git, Python 3.11.")

section_heading(doc, "C-2  Quick Start")
qs_code = (
    "# Clone repository\n"
    "git clone https://github.com/FBI-ZEZO03/NeuroTrap.git\n"
    "cd neurotrap\n\n"
    "# Configure environment\n"
    "cp .env.example .env   # edit with secure values\n\n"
    "# Generate SSL certificate\n"
    "bash scripts/generate_ssl_cert.sh\n\n"
    "# Configure UFW (honeypot + management ports)\n"
    "sudo ufw allow 22/tcp 23/tcp 21/tcp 80/tcp 443/tcp\n"
    "sudo ufw allow 445/tcp 1433/tcp 3306/tcp 3389/tcp 5900/tcp\n"
    "sudo ufw allow 8088/tcp 161/udp 50402/tcp\n"
    "sudo ufw enable\n\n"
    "# Build and start\n"
    "docker compose up -d --build\n\n"
    "# Initialize database indexes\n"
    "docker compose exec api python -m src.database.setup_db_indexes\n\n"
    "# Verify\n"
    "docker ps --format 'table {{.Names}}\\t{{.Status}}\\t{{.Ports}}'"
)
p = doc.add_paragraph()
run = p.add_run(qs_code)
run.font.name = 'Courier New'
run.font.size = Pt(10)
p.paragraph_format.space_after = Pt(12)

para(doc, "Page C2", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.RIGHT,
     space_before=0, space_after=12)

section_heading(doc, "C-3  Accessing the Dashboard")
body(doc,
     "Access https://[SERVER_IP] in a browser. Accept the self-signed certificate "
     "security warning. Log in with the admin credentials from .env. "
     "Default credentials (admin/neurotrap2024) MUST be changed before production "
     "deployment by updating ADMIN_PASS in .env and restarting the API container.")

section_heading(doc, "C-4  Enabling Optional Features")
optionals = [
    ("Galah LLM honeypot",
     "Add ANTHROPIC_API_KEY=sk-ant-... to .env, then: docker compose up -d galah"),
    ("LLM-powered SOC reports",
     "Add ANTHROPIC_API_KEY or GROQ_API_KEY to .env. SOC Analyst auto-switches to LLM mode."),
    ("MFA for admin login",
     "Generate secret: python -c \"import pyotp; print(pyotp.random_base32())\". "
     "Add MFA_ENABLED=1 and MFA_SECRET=[secret] to .env. Restart API container. "
     "Scan QR from /api/auth/otp/setup with authenticator app."),
]
for title, desc in optionals:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run1 = p.add_run(title + ": ")
    run1.bold = True; run1.font.name='Times New Roman'; run1.font.size=Pt(12)
    run2 = p.add_run(desc)
    run2.font.name='Times New Roman'; run2.font.size=Pt(12)

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
doc.save(OUTPUT_PATH)
print(f"Document saved: {OUTPUT_PATH}")
print(f"Sections: {len(doc.sections)}")
print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Tables: {len(doc.tables)}")
